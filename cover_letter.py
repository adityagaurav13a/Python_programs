import os
from docx import Document

def read_txt_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def write_txt_file(file_path, content):
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(content)

def read_docx_file(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def write_docx_file(file_path, content):
    doc = Document()
    for line in content.split('\n'):
        doc.add_paragraph(line)
    doc.save(file_path)

def replace_keywords(content, position_text, company_name_text):
    content = content.replace("position", position_text)
    content = content.replace("company_name", company_name_text)
    return content

def main():
    file_path = input("Enter the full path of the file (.txt or .docx): ").strip()
    if not os.path.exists(file_path):
        print("File does not exist.")
        return

    position_text = input("Enter the Job Position to replace: ")
    company_name_text = input("Enter the Company Name to replace: ")

    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".txt":
        content = read_txt_file(file_path)
    elif ext == ".docx":
        content = read_docx_file(file_path)
    else:
        print("Unsupported file format. Use .txt or .docx only.")
        return

    updated_content = replace_keywords(content, position_text, company_name_text)

    new_file_path = file_path.replace(ext, f"_updated{ext}")
    if ext == ".txt":
        write_txt_file(new_file_path, updated_content)
    elif ext == ".docx":
        write_docx_file(new_file_path, updated_content)

    print(f"Updated file saved as: {new_file_path}")

if __name__ == "__main__":
    main()
