import os
from docx import Document

def read_txt_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def write_txt_file(file_path, content):
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(content)

def read_docx_file(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def write_docx_file(file_path, content):
    doc = Document()
    for line in content.split('\n'):
        doc.add_paragraph(line)
    doc.save(file_path)

def replace_keywords(content, Job_title, Company_Name, Date, LinkedIn_URL, Contact, **kwargs):
    content = content.replace("{Job_title}", Job_title)
    content = content.replace("{Company_Name}", Company_Name)
    content = content.replace("{Date}", Date)
    content = content.replace("{LinkedIn_URL}", LinkedIn_URL)
    content = content.replace("{Contact}", Contact)
    return content

def main():
    file_path = input("Enter the full path of the file (.txt or .docx): ").strip()
    if not os.path.exists(file_path):
        print("File does not exist.")
        return

    Job_title = input("Enter the Job Position to replace: ")
    Company_Name = input("Enter the Company Name to replace: ")
    Date = input("Enter the Current Date (e.g., 2023-10-01): ")
    LinkedIn_URL = input("Enter your LinkedIn URL: ")
    Contact = input("Enter your Contact Information: ")
    
    # # Debugging line to set a breakpoint
    # import pdb
    # pdb.set_trace()

    # print("job title : " + Job_title)
    # print("company name : " + Company_Name)
    # print("date : " + Date)
    # print("LinkedIn URL : " + LinkedIn_URL)
    # print("contact : " + Contact)
    # print(" " )
    # print(" " )
    # print(" " )
    # print(" " )

    ext = os.path.splitext(file_path)[1].lower()
    # print("Extension :" +ext)
    
    if ext == ".txt":
        content = read_txt_file(file_path)
    elif ext == ".docx":
        content = read_docx_file(file_path)
    else:
        print("Unsupported file format. Use .txt or .docx only.")
        return

    updated_content = replace_keywords(content, Job_title, Company_Name, Date, LinkedIn_URL, Contact)

    new_file_path = file_path.replace(ext, f"_updated{ext}")
    if ext == ".txt":
        write_txt_file(new_file_path, updated_content)
    elif ext == ".docx":
        write_docx_file(new_file_path, updated_content)

    print(f"Updated file saved as: {new_file_path}")

if __name__ == "__main__":
    main()
